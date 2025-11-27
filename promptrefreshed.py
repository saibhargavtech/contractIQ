import os
import re
import json
import math
import itertools
from typing import List, Tuple
from collections import Counter, defaultdict
from urllib.parse import quote

import openai
import gradio as gr
import networkx as nx
import community as community_louvain
import matplotlib.pyplot as plt
import pandas as pd
from dotenv import load_dotenv

# ---------- Optional parsers ----------
PDF_AVAILABLE = True
DOCX_AVAILABLE = True
try:
    from PyPDF2 import PdfReader  # lightweight pdf text
except Exception:
    PDF_AVAILABLE = False

try:
    import docx  # python-docx
except Exception:
    DOCX_AVAILABLE = False

# ---------- Optional RDF ----------
RDFLIB_AVAILABLE = True
try:
    from rdflib import Graph as RDFGraph, Namespace, URIRef, BNode, Literal
    from rdflib.namespace import RDF, RDFS, OWL, SKOS, DCTERMS
except Exception:
    RDFLIB_AVAILABLE = False

# ---------- Optional LlamaIndex ----------
LLAMA_AVAILABLE = True
try:
    from llama_index.core import Document, VectorStoreIndex, Settings
    from llama_index.core.node_parser import SentenceSplitter
    from llama_index.embeddings.openai import OpenAIEmbedding
    from llama_index.llms.openai import OpenAI as LlamaOpenAI
except Exception:
    LLAMA_AVAILABLE = False

# ============ Setup ============
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT", "")
azure_api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview")
use_azure = bool(azure_endpoint)

if not api_key:
    raise ValueError("OPENAI_API_KEY environment variable is not set. Please set it in your .env file or environment.")

# Configure OpenAI client for Azure or standard OpenAI
if use_azure:
    # Extract deployment name and base URL from endpoint
    if "/deployments/" in azure_endpoint:
        parts = azure_endpoint.split("/deployments/")
        azure_base_url = parts[0]
        deployment_part = parts[1].split("/")[0]  # Get just the deployment name, before any /chat/completions
        if "?" in deployment_part:
            deployment_part = deployment_part.split("?")[0]
        azure_deployment = deployment_part
    else:
        azure_base_url = azure_endpoint
        azure_deployment = os.getenv("AZURE_DEPLOYMENT_NAME", "gpt-4.1")
    
    # Remove /openai from base URL if present (Azure OpenAI base URL should not include /openai)
    if azure_base_url.endswith("/openai"):
        azure_base_url = azure_base_url[:-7]
    
    # Set up Azure OpenAI
    openai.api_key = api_key
    openai.api_base = azure_base_url
    openai.api_type = "azure"
    openai.api_version = azure_api_version
    AZURE_DEPLOYMENT_NAME = azure_deployment
else:
    openai.api_key = api_key
    AZURE_DEPLOYMENT_NAME = None

# Globals
G = nx.DiGraph()
partition = {}
VECTOR_INDEX = None
ENABLE_VECTOR = False
VECTOR_TOPK = 8

SEARCH_OPTIONS = []
CLUSTER_SUMMARIES = {}  # {cluster_id: {"Cluster Summary":..., "KPI/Risk Label":...}}
GRAPH_CONTEXT_MEMORY = ""
RAW_COMBINED_TEXT = ""  # for LLM summaries

# ---- Q&A tuning (prompt compaction) ----
MAX_EDGES_IN_PROMPT = 48
MAX_EVIDENCE_CHARS = 220
EDGE_NAME_WEIGHT = 3
EDGE_REL_WEIGHT  = 2
EDGE_EVID_WEIGHT = 1

# ====== UI-Tunable Defaults (Graph Tuning Panel) ======
EXTRACT_TEMP = 0.4
STRICT_PROMPT = True
REQUIRE_EVIDENCE = True
ENABLE_FALLBACK_COOCC = False
MAX_COOCC_PAIRS = 5
CHUNK_MAX_CHARS = 3500

RELATION_POLICY = "Off"  # Off | Standard | Strict
ENABLE_ALIAS_NORMALIZATION = True
SKIP_EDGES_NO_EVID = True

LOUVAIN_RESOLUTION = 1.25
LOUVAIN_RANDOM_STATE = 202

# ====== File helpers (PDF/Word/Text only) ======
SUPPORTED_TYPES = {
    "PDF": [".pdf"],
    "Word": [".docx"],
    "Text": [".txt"]
}

def _normalize_types(selected: List[str]) -> List[str]:
    exts = []
    for t in selected or []:
        exts += SUPPORTED_TYPES.get(t, [])
    return list(dict.fromkeys(exts))

def _read_pdf(path: str) -> str:
    if not PDF_AVAILABLE:
        print(f"[PDF] PyPDF2 not available; skipping: {path}")
        return ""
    try:
        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(txt)
        return "\n".join(parts)
    except Exception as e:
        print(f"[PDF] Failed to parse {path}: {e}")
        return ""

def _read_docx(path: str) -> str:
    if not DOCX_AVAILABLE:
        print(f"[DOCX] python-docx not available; skipping: {path}")
        return ""
    try:
        d = docx.Document(path)
        paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        return "\n".join(paras)
    except Exception as e:
        print(f"[DOCX] Failed to parse {path}: {e}")
        return ""

def _read_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"[TXT] Failed to read {path}: {e}")
        return ""

def _read_one_file(path: str) -> str:
    p = path.lower()
    if p.endswith(".pdf"):
        return _read_pdf(path)
    if p.endswith(".docx"):
        return _read_docx(path)
    if p.endswith(".txt"):
        return _read_txt(path)
    return ""

def _list_dir_files(root: str, allowed_exts: List[str]) -> List[str]:
    out = []
    if not root or not os.path.isdir(root):
        return out
    for dirpath, _dirs, files in os.walk(root):
        for fn in files:
            ext = os.path.splitext(fn)[1].lower()
            if allowed_exts and ext not in allowed_exts:
                continue
            out.append(os.path.join(dirpath, fn))
    return out

# ====== Utility helpers ======
def _entropy(values):
    if not values: return 0.0
    total = sum(values)
    if total <= 0: return 0.0
    ps = [v / total for v in values if v > 0]
    return -sum(p * math.log(p, 2) for p in ps)

def _gini(seq):
    xs = sorted(max(0.0, float(x)) for x in seq)
    n = len(xs)
    if n == 0: return 0.0
    s = sum(xs)
    if s == 0: return 0.0
    num = 0.0
    for i, x in enumerate(xs, start=1):
        num += (2*i - n - 1) * x
    return num / (n * s)

def _giant_component_undirected(Gu):
    if Gu.number_of_nodes() == 0: return Gu.copy()
    comps = list(nx.connected_components(Gu))
    if not comps: return Gu.copy()
    largest = max(comps, key=len)
    return Gu.subgraph(largest).copy()

def _chunk(text, max_chars=None):
    max_chars = max_chars or CHUNK_MAX_CHARS
    buf, out = [], []
    for para in re.split(r"\n{2,}", text):
        if sum(len(x) for x in buf) + len(para) + 1 > max_chars:
            if buf:
                out.append("\n\n".join(buf)); buf = []
        if para.strip():
            buf.append(para.strip())
    if buf: out.append("\n\n".join(buf))
    return out

def _parse_json_or_lines(raw):
    triples = []
    raw = (raw or "").strip()
    try:
        data = json.loads(raw)
        if isinstance(data, dict) and "triples" in data:
            data = data["triples"]
        if isinstance(data, list):
            for t in data:
                h = (t.get("head") or t.get("subject") or "").strip()
                r = (t.get("relation") or t.get("predicate") or "").strip()
                ta = (t.get("tail") or t.get("object") or "").strip()
                ev = (t.get("evidence") or t.get("source") or "").strip()
                if h and r and ta:
                    triples.append((h, r, ta, ev))
            return triples
    except Exception:
        pass
    for line in raw.splitlines():
        if "--[" in line and "]-->" in line:
            try:
                lhs, rest = line.split("--[", 1)
                rel, rhs = rest.split("]-->", 1)
                h = lhs.strip(); r = rel.strip(); ta = rhs.strip()
                if h and r and ta:
                    triples.append((h, r, ta, ""))  # no evidence available
            except Exception:
                continue
    return triples

def _fallback_cooccurrence(sent):
    tokens = re.findall(r"\b[A-Z][A-Za-z0-9&\-/\.]{2,}\b", sent)
    tokens = [t for t in tokens if t not in {"The","And","For","With","Note","Page"}]
    tokens = list(dict.fromkeys(tokens))
    pairs = list(itertools.combinations(tokens, 2))
    if not ENABLE_FALLBACK_COOCC: return []
    if MAX_COOCC_PAIRS > 0: pairs = pairs[:MAX_COOCC_PAIRS]
    return [(a, "CO_OCCURS_IN_SENTENCE", b, sent.strip()) for a, b in pairs]

# ====== Relation policy / alias normalization ======
REL_ALLOW_STANDARD = {
    "governed_by","audited_by","recognised_under","modifies",
    "involves_counterparty","has_obligation","linked_to","evidenced_by"
}
REL_ALLOW_STRICT = {
    "governed_by","audited_by","recognised_under","has_obligation","evidenced_by"
}
REL_MAP = {"recognized_under":"recognised_under", "recognized by":"recognised_under", "governed by":"governed_by"}
ALIASES = {
    "asc 606":"ASC 606","ifrs 15":"IFRS 15","asc 842":"ASC 842","ifrs 16":"IFRS 16",
    "right-of-use asset":"ROU asset","acme holdings limited":"Acme Holdings Ltd."
}
def _norm_rel(rel):
    rel = (rel or "").strip().lower().replace(" ", "_")
    return REL_MAP.get(rel, rel)

def _relation_allowed(rel_norm):
    if RELATION_POLICY == "Off": return True
    if RELATION_POLICY == "Standard": return rel_norm in REL_ALLOW_STANDARD
    if RELATION_POLICY == "Strict": return rel_norm in REL_ALLOW_STRICT
    return True

def _canon_entity(name):
    if not ENABLE_ALIAS_NORMALIZATION: return name.strip()
    return ALIASES.get(name.strip().lower(), name.strip())

# ====== Resilient LLM call ======
def _chat_strict(messages, temperature=0.2, timeout=45, model="gpt-4o"):
    # Works for both Azure OpenAI and standard OpenAI; falls back gracefully
    try:
        if use_azure:
            # Use Azure OpenAI
            from openai import AzureOpenAI
            client = AzureOpenAI(
                api_key=api_key,
                api_version=azure_api_version,
                azure_endpoint=azure_base_url
            )
            # Use deployment name instead of model name for Azure
            comp = client.chat.completions.create(
                model=AZURE_DEPLOYMENT_NAME,
                messages=messages,
                temperature=float(temperature),
                timeout=timeout
            )
        else:
            # Use standard OpenAI
            comp = openai.chat.completions.create(
                model=model, messages=messages, temperature=float(temperature), timeout=timeout
            )
        return (comp.choices[0].message.content or "").strip()
    except Exception as e:
        print(f"[⚠️ LLM ERROR] {str(e)}")
        pass
    try:
        if use_azure:
            # Legacy Azure OpenAI format
            comp = openai.ChatCompletion.create(
                engine=AZURE_DEPLOYMENT_NAME,
                messages=messages,
                temperature=float(temperature),
                timeout=timeout
            )
        else:
            # Legacy standard OpenAI format
            comp = openai.ChatCompletion.create(
                model=model, messages=messages, temperature=float(temperature), timeout=timeout
            )
        return (comp.choices[0].message.content or "").strip()
    except Exception as e:
        print(f"[⚠️ LLM ERROR] {str(e)}")
        pass
    
    # Ultimate fallback - return empty result to prevent hanging
    print(f"[⚠️ LLM TIMEOUT] OpenAI call timed out after {timeout}s - skipping LLM extraction")
    return ""

# ====== Simple rule-based extractor (fallback) ======
_REL_PATTERNS = [
    (re.compile(r"(.+?)\s+(?:is|was|are|were)?\s*audited by\s+(.+?)\.", re.I), "audited_by"),
    (re.compile(r"(.+?)\s+(?:is|was|are|were)?\s*governed by\s+(.+?)\.", re.I), "governed_by"),
    (re.compile(r"(.+?)\s+(?:is|was|are|were)?\s*recognis(?:e|z)d? under\s+(.+?)\.", re.I), "recognised_under"),
    (re.compile(r"(.+?)\s+has obligation(?:s)? to\s+(.+?)\.", re.I), "has_obligation"),
    (re.compile(r"(.+?)\s+involves\s+(.+?)\.", re.I), "involves_counterparty"),
]
def _rule_based_triples(text):
    triples = []
    for sent in re.split(r'(?<=[.!?])\s+', text):
        s = sent.strip()
        if not s: continue
        for pat, rel in _REL_PATTERNS:
            m = pat.search(s)
            if m:
                a = m.group(1).strip().strip('",() ')
                b = m.group(2).strip().strip('",() ')
                if a and b and a.lower() != b.lower():
                    triples.append((a, rel, b, s))
        kw_hits = [
            ("recognised_under", ["ifrs", "asc", "gaap", "standard", "policy"]),
            ("audited_by", ["auditor", "audit", "audited"]),
            ("governed_by", ["governed", "regulation", "law", "act"]),
        ]
        low = s.lower()
        for rel, kws in kw_hits:
            if any(k in low for k in kws):
                ents = re.findall(r"\b[A-Z][A-Za-z0-9&\-/\.]{2,}(?:\s+[A-Z][A-Za-z0-9&\-/\.]{2,})*\b", s)
                if len(ents) >= 2:
                    triples.append((ents[0], rel, ents[1], s))
    out, seen = [], set()
    for a, r, b, ev in triples:
        key = (a, r, b, ev)
        if key not in seen:
            seen.add(key); out.append(key)
    return out

# ============ 1) Entity & Relation Extraction ============
def extract_entities_relations(text: str) -> List[Tuple[str, str, str, str]]:
    content = text[:20000]
    chunks = _chunk(content, max_chars=CHUNK_MAX_CHARS)
    triples_with_source = []
    
    # Quick bypass for faster processing - use rule-based extraction only
    print(f"[🚀 FAST MODE] Using rule-based extraction for faster processing...")
    triples_with_source = _rule_based_triples(content)
    print(f"[✅ FAST MODE] Rule-based extraction found {len(triples_with_source)} triples")
    return triples_with_source

    base_prompt = (
        'Extract financial entity–relation–entity triples from the text. '
        'Return STRICT JSON as: '
        '{"triples":[{"head":"...", "relation":"...", "tail":"...", "evidence":"original sentence"}]}.\n'
        'Guidelines:\n'
        '- Use short, canonical entity names.\n'
        '- Relations must be concise verbs/phrases (e.g., "recognised_under","audited_by","governed_by").\n'
        '- "evidence" must be the exact sentence containing the relation.\n'
        '- Skip trivial or ambiguous lines.\n'
    )
    if STRICT_PROMPT:
        base_prompt += (
            '- Reject edges if the relation verb is missing.\n'
            '- If evidence is unavailable, return an empty "triples" array.\n'
            '- Do not infer beyond text; avoid co-occurrence-only edges.\n'
        )
    base_prompt += 'Text:\n'

    for ch_idx, ch in enumerate(chunks):
        print(f"[🔍 LLM EXTRACTION] Processing chunk {ch_idx+1}/{len(chunks)} ({len(ch)} chars)...")
        msg = [{"role": "user", "content": base_prompt + ch}]
        try:
            # Add aggressive timeout - if LLM hangs, skip it
            raw = _chat_strict(msg, temperature=EXTRACT_TEMP, timeout=10, model="gpt-4o")
            if raw:  # Only process if we got a response
                parsed = _parse_json_or_lines(raw)
                triples_with_source.extend(parsed)
                print(f"[✅ LLM EXTRACTION] Chunk {ch_idx+1} extracted {len(parsed)} triples")
            else:
                print(f"[⚠️ LLM EXTRACTION] Chunk {ch_idx+1} timed out - skipping to fallback")
        except Exception as e:
            print(f"[❌ LLM EXTRACTION] Chunk {ch_idx+1} failed: {e}")
            pass

    if not triples_with_source:
        print(f"[🔄 FALLBACK] LLM extraction failed - using rule-based extraction...")
        triples_with_source = _rule_based_triples(content)
        print(f"[✅ FALLBACK] Rule-based extraction found {len(triples_with_source)} triples")

    if not triples_with_source and ENABLE_FALLBACK_COOCC:
        print(f"[🔄 COOCCURRENCE] Rule-based failed - using co-occurrence fallback...")
        for s in re.split(r'(?<=[.!?])\s+', content):
            if s.strip():
                triples_with_source.extend(_fallback_cooccurrence(s))
        print(f"[✅ COOCCURRENCE] Co-occurrence fallback found {len(triples_with_source)} triples")

    cleaned, seen = [], set()
    for a, r, b, ev in triples_with_source:
        a = _canon_entity(a); b = _canon_entity(b)
        r_norm = _norm_rel(r)
        ev = (ev or "").strip()
        if not a or not b or a.lower() == b.lower(): continue
        if not _relation_allowed(r_norm): continue
        if (SKIP_EDGES_NO_EVID or REQUIRE_EVIDENCE) and not ev:
            continue
        key = (a, r_norm, b, ev)
        if key in seen: continue
        seen.add(key); cleaned.append(key)
    return cleaned

# ============ 2) Build Graph ============
def build_graph(extraction_output, source="document"):
    global G
    if G is None or not isinstance(G, nx.DiGraph):
        G = nx.DiGraph()
    for a, rel, b, evidence in extraction_output:
        if G.has_edge(a, b):
            continue
        G.add_edge(a, b, label=rel, source=source, evidence_text=evidence)
    return G

# ============ 2b) Multi-file pipeline glue ============
def load_corpus(upload_paths: List[str], dir_path: str, type_choices: List[str]) -> List[Tuple[str, str]]:
    """
    Returns list of (source_name, text)
    - upload_paths: list of paths from Gradio uploader (can be None)
    - dir_path: a directory path (can be empty)
    - type_choices: ["PDF","Word","Text"] -> filters extensions
    """
    allowed_exts = _normalize_types(type_choices)
    pairs = []

    # uploads
    for p in (upload_paths or []):
        if not p: continue
        ext = os.path.splitext(str(p))[1].lower()
        if allowed_exts and ext not in allowed_exts:
            continue
        txt = _read_one_file(str(p))
        if txt.strip():
                pairs.append((os.path.basename(str(p)), txt))

    # directory
    if dir_path and os.path.isdir(dir_path):
        for p in _list_dir_files(dir_path, allowed_exts):
            txt = _read_one_file(p)
            if txt.strip():
                pairs.append((os.path.relpath(p, dir_path), txt))

    return pairs

# ============ 3) Community Detection ============
ONTO_REL_BOOST = {
    "recognised_under": 1.50,
    "audited_by": 1.40,
    "governed_by": 1.30,
    "has_obligation": 1.20,
    "involves_counterparty": 1.10,
}

def _to_weighted_undirected(G: nx.DiGraph, use_ontology_for_clustering: bool) -> nx.Graph:
    Gu = nx.Graph()
    for u, v, d in G.edges(data=True):
        base = 1.0
        if use_ontology_for_clustering:
            rel = (d.get("label") or "").lower()
            base *= ONTO_REL_BOOST.get(rel, 1.0)
        if Gu.has_edge(u, v):
            Gu[u][v]["weight"] += base
        else:
            Gu.add_edge(u, v, weight=base)
    return Gu

def detect_communities(G, resolution=None, random_state=None, use_ontology_for_clustering=False):
    if G.number_of_nodes() == 0:
        return G, {}
    res = resolution if resolution is not None else LOUVAIN_RESOLUTION
    rs  = random_state if random_state is not None else LOUVAIN_RANDOM_STATE
    Gu = _to_weighted_undirected(G, use_ontology_for_clustering)
    part = community_louvain.best_partition(Gu, resolution=float(res), random_state=int(rs), weight="weight")
    nx.set_node_attributes(G, part, 'community')
    return G, part

# ============ 4) Cluster Summarization ============
def _heuristic_cluster_label_and_summary(G, cid, nodes, full_text):
    sub = G.subgraph(nodes).copy()
    degrees = dict(sub.degree())
    top_nodes = [n for n, _ in sorted(degrees.items(), key=lambda x: x[1], reverse=True)[:3]]
    top_nodes_str = ", ".join(top_nodes) if top_nodes else ", ".join(nodes[:3])
    label = "Theme / Risk"
    summary = (
        f"This cluster groups {len(nodes)} entities with dense links among {top_nodes_str}. "
        f"Theme inferred from connectivity and relation evidence."
    )
    return label, summary

def summarize_clusters(G, part, full_text):
    if not part: return {}
    communities = {}
    for node, group in part.items():
        communities.setdefault(group, []).append(node)

    combined = {}
    context = (full_text or "")[:2000]

    for comm_id, nodes in communities.items():
        ent_list = ', '.join(nodes[:50])
        fb_label, fb_summary = _heuristic_cluster_label_and_summary(G, comm_id, nodes, context)
        summary = fb_summary; label = fb_label

        # Best-effort LLM
        try:
            if use_azure:
                from openai import AzureOpenAI
                client = AzureOpenAI(
                    api_key=api_key,
                    api_version=azure_api_version,
                    azure_endpoint=azure_base_url
                )
                summary_response = client.chat.completions.create(
                    model=AZURE_DEPLOYMENT_NAME,
                    messages=[{"role": "user", "content": f"""
You are a financial analyst assistant. Given the following entities: {ent_list},
and the original context:
\"\"\"{context}\"\"\" 
Summarize this group's theme or risk in 2–3 sentences.
"""}],
                    temperature=0.3, timeout=15
                )
            else:
                summary_response = openai.ChatCompletion.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": f"""
You are a financial analyst assistant. Given the following entities: {ent_list},
and the original context:
\"\"\"{context}\"\"\" 
Summarize this group's theme or risk in 2–3 sentences.
"""}],
                    temperature=0.3, timeout=15
                )
            _s = (summary_response.choices[0].message.content or "").strip()
            if _s: summary = _s
        except Exception:
            pass
        try:
            if use_azure:
                from openai import AzureOpenAI
                client = AzureOpenAI(
                    api_key=api_key,
                    api_version=azure_api_version,
                    azure_endpoint=azure_base_url
                )
                label_response = client.chat.completions.create(
                    model=AZURE_DEPLOYMENT_NAME,
                    messages=[{"role": "user", "content": f"""
Given entities: {ent_list} and the context above,
Suggest one concise KPI or risk label. Reply with only the label.
"""}],
                    temperature=0.2, timeout=10
                )
            else:
                label_response = openai.ChatCompletion.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": f"""
Given entities: {ent_list} and the context above,
Suggest one concise KPI or risk label. Reply with only the label.
"""}],
                    temperature=0.2, timeout=10
                )
            _l = (label_response.choices[0].message.content or "").strip()
            if _l: label = _l
        except Exception:
            pass

        combined[comm_id] = {
            "Cluster Summary": summary,
            "KPI/Risk Label": label,
        }
    return combined

# ============ 4.5) Contract Context Creation ============
def create_contract_context_text(contract_df):
    """
    Convert contract DataFrame to rich text context for GraphRAG
    This enriches the knowledge graph with structured financial and business data
    """
    if contract_df is None or contract_df.empty:
        return ""
    
    context_parts = ["CONTRACT PORTFOLIO BUSINESS INTELLIGENCE:"]
    
    # Portfolio-level insights
    context_parts.append(f"The organization has {len(contract_df)} active contracts.")
    context_parts.append(f"Total contract value: ${contract_df['total_value_usd'].sum():,.0f}")
    context_parts.append(f"Total annual commitment: ${contract_df['annual_commitment_usd'].sum():,.0f}")
    
    # Vendor analysis
    vendor_counts = contract_df['counterparty'].value_counts()
    top_vendors = vendor_counts.head(5)
    context_parts.append("Top vendors by contract volume:")
    for vendor, count in top_vendors.items():
        vendor_value = contract_df[contract_df['counterparty'] == vendor]['total_value_usd'].sum()
        context_parts.append(f"- {vendor}: {count} contracts, ${vendor_value:,.0f} total value")
    
    # Contract types analysis  
    type_summary = contract_df['type'].value_counts()
    context_parts.append("Contract distribution by type:")
    for contract_type, count in type_summary.items():
        context_parts.append(f"- {contract_type}: {count} contracts")
    
    # Financial risk analysis
    high_value_contracts = contract_df[contract_df['total_value_usd'] > 1000000]
    if not high_value_contracts.empty:
        context_parts.append(f"High-value contracts (>$1M): {len(high_value_contracts)} contracts totaling ${high_value_contracts['total_value_usd'].sum():,.0f}")
    
    # Compliance and risk insights
    compliance_summary = contract_df['compliance'].value_counts()
    context_parts.append("Compliance framework coverage:")
    for compliance, count in compliance_summary.items():
        context_parts.append(f"- {compliance}: {count} contracts")
    
    # SLA performance
    sla_summary = contract_df['sla_uptime'].value_counts()
    context_parts.append("Service level expectations:")
    for sla, count in sla_summary.items():
        context_parts.append(f"- {sla} uptime: {count} contracts")
    
    # Contract summaries for additional context
    if 'summary' in contract_df.columns:
        context_parts.append("\nINDIVIDUAL CONTRACT SUMMARIES:")
        for idx, row in contract_df.iterrows():
            if pd.notna(row.get('summary')):
                context_parts.append(f"\nContract {row['contract_id']} ({row['counterparty']}):")
                context_parts.append(row['summary'])
    
    return "\n".join(context_parts)

# ============ 4.6) Graph Context Memory ============
def build_graph_context_memory(G, part, summaries, max_edges_per_cluster=6, max_nodes_list=8):
    if not part: return "(no graph memory)"
    communities = defaultdict(list)
    for node, cid in part.items():
        communities[cid].append(node)

    lines = ["GRAPH CONTEXT MEMORY"]
    for cid, nodes in sorted(communities.items(), key=lambda x: x[0]):
        sub = G.subgraph(nodes).copy()
        deg = dict(sub.degree())
        top_nodes = [n for n, _ in sorted(deg.items(), key=lambda x: x[1], reverse=True)[:max_nodes_list]]
        s = summaries.get(cid, {})
        label = s.get("KPI/Risk Label", "")
        summ  = s.get("Cluster Summary", "")

        lines.append(f"\n[Cluster {cid}] Label: {label}")
        lines.append(f"[Cluster {cid}] Summary: {summ}")
        if top_nodes:
            lines.append(f"[Cluster {cid}] Key entities: {', '.join(top_nodes)}")

        k = 0
        for u, v, d in sub.edges(data=True):
            if k >= max_edges_per_cluster: break
            lbl = d.get("label",""); ev  = d.get("evidence_text","")
            if lbl or ev:
                if len(ev) > MAX_EVIDENCE_CHARS: ev = ev[:MAX_EVIDENCE_CHARS].rstrip() + "…"
                lines.append(f"[Cluster {cid}] Rel: {u} —[{lbl}]→ {v}; Ev: {ev}")
                k += 1
    return "\n".join(lines)

# ============ 5) Visualize Graph ============
def visualize_graph(G):
    plt.figure(figsize=(12, 8))
    if G.number_of_nodes() == 0:
        plt.title("GraphRAG Knowledge Graph (no edges extracted)")
        plt.axis('off')
        try: 
            plt.tight_layout(pad=1.0)
        except Exception as e:
            print(f"[Warning] Empty graph layout failed: {e}")
            plt.subplots_adjust(left=0.1, right=0.9, top=0.9, bottom=0.1)
        plt.savefig("graphrag_output.png", dpi=150, bbox_inches='tight'); plt.close()
        return "graphrag_output.png"

    pos = nx.spring_layout(G, seed=42, k=0.7)
    colors = [G.nodes[n].get('community', 0) for n in G.nodes()]
    nx.draw(
        G, pos, with_labels=True,
        node_color=colors, cmap=plt.cm.Set3,
        edge_color="#999999", font_size=9, node_size=900,
        arrowsize=12, width=1.2
    )
    edge_labels = {(u, v): d.get('label', '') for u, v, d in G.edges(data=True)}
    nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, font_size=8, label_pos=0.5)
    plt.title("GraphRAG Knowledge Graph")
    plt.axis('off')
    try: 
        plt.tight_layout(pad=1.0)
    except Exception as e:
        print(f"[Warning] Layout adjustment failed: {e}")
        plt.subplots_adjust(left=0.1, right=0.9, top=0.9, bottom=0.1)
    plt.savefig("graphrag_output.png", dpi=150, bbox_inches='tight'); plt.close()
    return "graphrag_output.png"

# ============ 6) Metrics ============
def compute_metrics(G, part):
    V = G.number_of_nodes(); E = G.number_of_edges()
    density = nx.density(G)
    avg_out = E / V if V > 0 else 0.0
    avg_in  = E / V if V > 0 else 0.0

    Gu = G.to_undirected()
    degs = [d for _, d in Gu.degree()]
    deg_hist = Counter(degs)
    entropy = _entropy(list(deg_hist.values()))
    gini = _gini(degs)

    try:
        assort = nx.degree_assortativity_coefficient(Gu) if Gu.number_of_edges() > 0 and len(set(degs)) > 1 else 0.0
    except (ZeroDivisionError, ValueError, Exception) as e:
        print(f"[Warning] Assortativity calculation failed: {e}")
        assort = 0.0
    try:
        clustering = nx.transitivity(Gu) if Gu.number_of_edges() > 0 else 0.0
    except Exception as e:
        print(f"[Warning] Clustering calculation failed: {e}")
        clustering = 0.0

    wcc_count = nx.number_weakly_connected_components(G) if V > 0 else 0
    giant_ratio = 0.0
    if V > 0 and wcc_count > 0:
        largest = max(nx.weakly_connected_components(G), key=len)
        giant_ratio = len(largest) / V if V > 0 else 0.0

    asp = 0.0; diam = 0
    if Gu.number_of_nodes() > 1:
        Gc = _giant_component_undirected(Gu)
        try:
            if Gc.number_of_nodes() > 1:
                asp = nx.average_shortest_path_length(Gc)
                diam = nx.diameter(Gc)
        except (nx.NetworkXError, nx.NetworkXNoPath, Exception) as e:
            print(f"[Warning] Path calculation failed: {e}")
            asp, diam = 0.0, 0

    modularity = 0.0
    try:
        if part:
            modularity = community_louvain.modularity(part, Gu)
    except Exception:
        modularity = 0.0

    edges_with_evidence = sum(1 for _, _, d in G.edges(data=True) if d.get('evidence_text'))
    prov_completeness = (edges_with_evidence / E) if E > 0 else 0.0

    rel_types = set(d.get('label') for _, _, d in G.edges(data=True) if d.get('label'))
    rel_type_count = len(rel_types)

    r = lambda x: round(float(x), 3)
    rows = [
        ("Node count", V),
        ("Edge count", E),
        ("Density (directed)", r(density)),
        ("Avg out-degree", r(avg_out)),
        ("Avg in-degree", r(avg_in)),
        ("Degree entropy (undirected)", r(entropy)),
        ("Degree Gini (undirected)", r(gini)),
        ("Assortativity (undirected degree)", r(assort)),
        ("Global clustering coeff.", r(clustering)),
        ("Weakly connected components", int(wcc_count)),
        ("Giant component ratio", r(giant_ratio)),
        ("Avg shortest path (giant, undirected)", r(asp)),
        ("Diameter (giant, undirected)", int(diam)),
        ("Modularity (Louvain on undirected)", r(modularity)),
        ("Provenance completeness (edges)", r(prov_completeness)),
        ("Unique relation types", int(rel_type_count)),
    ]
    df = pd.DataFrame(rows, columns=["Metric", "Value"])
    return df

# === LlamaIndex: build & query vector index (fixed for multi-doc) ===
def build_vector_index_from_docs(doc_pairs: List[Tuple[str, str]]):
    """
    doc_pairs: list of (source_name, text)
    Builds a multi-document index with sensible chunking.
    """
    global VECTOR_INDEX
    VECTOR_INDEX = None
    if not (LLAMA_AVAILABLE and doc_pairs):
        return None

    try:
        if use_azure:
            # Azure OpenAI for LlamaIndex
            from llama_index.llms.azure_openai import AzureOpenAI
            from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
            llm = AzureOpenAI(
                model=AZURE_DEPLOYMENT_NAME,
                deployment_name=AZURE_DEPLOYMENT_NAME,
                api_key=api_key,
                azure_endpoint=azure_base_url,
                api_version=azure_api_version,
                temperature=0.0
            )
            embed_model = AzureOpenAIEmbedding(
                model="text-embedding-3-small",
                deployment_name="text-embedding-3-small",
                api_key=api_key,
                azure_endpoint=azure_base_url,
                api_version=azure_api_version
            )
        else:
            llm = LlamaOpenAI(model="gpt-4o", temperature=0.0, api_key=openai.api_key)
            embed_model = OpenAIEmbedding(model="text-embedding-3-small", api_key=openai.api_key)
        Settings.llm = llm
        Settings.embed_model = embed_model
    except Exception as e:
        print(f"[LlamaIndex] Failed to init OpenAI providers: {e}")
        return None

    splitter = SentenceSplitter(chunk_size=900, chunk_overlap=120)
    documents = []
    for src, txt in doc_pairs:
        if not txt.strip():
            continue
        doc = Document(text=txt, doc_id=src, metadata={"source": src})
        nodes = splitter.get_nodes_from_documents([doc])
        for n in nodes:
            node_text = n.get_content() if hasattr(n, "get_content") else getattr(n, "text", "")
            if node_text and node_text.strip():
                documents.append(Document(text=node_text, metadata={"source": src}))

    if not documents:
        return None

    try:
        VECTOR_INDEX = VectorStoreIndex.from_documents(documents)
    except Exception as e:
        print(f"[LlamaIndex] Index build failed: {e}")
        VECTOR_INDEX = None

    return VECTOR_INDEX

def vector_snippets(query, top_k=8):
    if not LLAMA_AVAILABLE or VECTOR_INDEX is None or not query:
        return []
    try:
        retriever = VECTOR_INDEX.as_retriever(similarity_top_k=int(top_k))
        results = retriever.retrieve(query)
        snippets = []
        for r in results:
            try:
                txt = r.node.get_content()
            except Exception:
                txt = getattr(r, "text", "")
            meta = r.node.metadata if hasattr(r.node, "metadata") else {}
            src = meta.get("source", "uploaded")
            if txt:
                cleaned = re.sub(r"\s+", " ", txt).strip()
                if len(cleaned) > 300:
                    cleaned = cleaned[:297] + "..."
                snippets.append(f"[{src}] {cleaned}")
        return snippets
    except Exception as e:
        print(f"[LlamaIndex] Retrieval failed: {e}")
        return []

# ============ 7) Full Pipeline ============
def _build_search_options(G, part):
    choices = []
    cids = sorted(set(part.values())) if part else []
    choices += [f"Cluster {cid}" for cid in cids]
    choices += sorted([str(n) for n in G.nodes()])
    return choices

# ============ 7.1) Contract Data Extraction Integration ============
def process_corpus_with_contract_extraction(
    upload_paths,
    dir_path,
    type_choices,
    enable_vector=False,
    vector_topk=8,
    extract_temp=EXTRACT_TEMP,
    strict_prompt=STRICT_PROMPT,
    require_evidence=REQUIRE_EVIDENCE,
    enable_fallback=ENABLE_FALLBACK_COOCC,
    max_coocc_pairs=MAX_COOCC_PAIRS,
    chunk_size=CHUNK_MAX_CHARS,
    relation_policy=RELATION_POLICY,
    alias_norm=ENABLE_ALIAS_NORMALIZATION,
    skip_no_evidence=SKIP_EDGES_NO_EVID,
    louvain_resolution=LOUVAIN_RESOLUTION,
    louvain_random_state=LOUVAIN_RANDOM_STATE,
    use_onto_cluster_cb=False,
):
    """Enhanced processing with error handling"""
    try:
        return _process_corpus_with_contract_extraction(
            upload_paths, dir_path, type_choices, enable_vector, vector_topk,
            extract_temp, strict_prompt, require_evidence, enable_fallback,
            max_coocc_pairs, chunk_size, relation_policy, alias_norm,
            skip_no_evidence, louvain_resolution, louvain_random_state,
            use_onto_cluster_cb
        )
    except Exception as e:
        print(f"[Error] Processing failed: {e}")
        # Return safe default outputs
        blank_img = visualize_graph(nx.DiGraph())
        info_df = pd.DataFrame([["Processing failed. Please try again."]], columns=["Error"])
        return blank_img, pd.DataFrame(), info_df, gr.update(choices=[], value=None), None, None, None, None, f"❌ Error: {str(e)}"

def _process_corpus_with_contract_extraction(
    upload_paths,
    dir_path,
    type_choices,
    enable_vector=False,
    vector_topk=8,
    extract_temp=EXTRACT_TEMP,
    strict_prompt=STRICT_PROMPT,
    require_evidence=REQUIRE_EVIDENCE,
    enable_fallback=ENABLE_FALLBACK_COOCC,
    max_coocc_pairs=MAX_COOCC_PAIRS,
    chunk_size=CHUNK_MAX_CHARS,
    relation_policy=RELATION_POLICY,
    alias_norm=ENABLE_ALIAS_NORMALIZATION,
    skip_no_evidence=SKIP_EDGES_NO_EVID,
    louvain_resolution=LOUVAIN_RESOLUTION,
    louvain_random_state=LOUVAIN_RANDOM_STATE,
    use_onto_cluster_cb=False,
):
    """
    Enhanced processing that includes contract data extraction for dashboard integration
    """
    global G, partition, ENABLE_VECTOR, VECTOR_TOPK, SEARCH_OPTIONS, CLUSTER_SUMMARIES, GRAPH_CONTEXT_MEMORY, RAW_COMBINED_TEXT
    global EXTRACT_TEMP, STRICT_PROMPT, REQUIRE_EVIDENCE, ENABLE_FALLBACK_COOCC, MAX_COOCC_PAIRS, CHUNK_MAX_CHARS
    global RELATION_POLICY, ENABLE_ALIAS_NORMALIZATION, SKIP_EDGES_NO_EVID, LOUVAIN_RESOLUTION, LOUVAIN_RANDOM_STATE

    ENABLE_VECTOR = bool(enable_vector)
    VECTOR_TOPK = int(vector_topk)
    EXTRACT_TEMP = float(extract_temp)
    STRICT_PROMPT = bool(strict_prompt)
    REQUIRE_EVIDENCE = bool(require_evidence)
    ENABLE_FALLBACK_COOCC = bool(enable_fallback)
    MAX_COOCC_PAIRS = int(max_coocc_pairs)
    CHUNK_MAX_CHARS = int(chunk_size)
    RELATION_POLICY = relation_policy
    ENABLE_ALIAS_NORMALIZATION = bool(alias_norm)
    SKIP_EDGES_NO_EVID = bool(skip_no_evidence)
    LOUVAIN_RESOLUTION = float(louvain_resolution)
    LOUVAIN_RANDOM_STATE = int(louvain_random_state)

    # 1) Load documents
    print(f"[📁 DOCUMENT LOADING] Starting with: upload_paths={upload_paths}, dir_path={dir_path}, type_choices={type_choices}")
    states = load_corpus(upload_paths, dir_path, type_choices)
    print(f"[📁 DOCUMENT LOADING] Loaded {len(states)} document pairs")
    for i, (src, txt) in enumerate(states):
        print(f"[📁 DOCUMENT {i+1}] {src}: {len(txt)} characters extracted")
    
    if not states:
        blank_img = visualize_graph(nx.DiGraph())
        info_df = pd.DataFrame([["No files were found/parsed. Check your directory path and type filters."]], columns=["Info"])
        SEARCH_OPTIONS = []; GRAPH_CONTEXT_MEMORY = "(no graph memory)"; RAW_COMBINED_TEXT = ""
        print("[❌ DOCUMENT LOADING] No files processed")
        return blank_img, pd.DataFrame(), info_df, gr.update(choices=[], value=None), None, None, None, None, "No files processed"

    RAW_COMBINED_TEXT = "\n\n".join([f"### {src}\n{txt}" for (src, txt) in states])

    # 2) Optional vector index
    if ENABLE_VECTOR:
        build_vector_index_from_docs(states)
    
    # 4) Extract & build graph (EXACTLY like primary.py - ONLY uploaded documents)
    G = nx.DiGraph()
    total_triples = 0
    print(f"[🧠 GRAPH BUILDING] Processing {len(states)} uploaded documents...")
    
    for src, txt in states:
        print(f"[📄 PROCESSING] {src} ({len(txt)} chars)")
        extracted = extract_entities_relations(txt)
        print(f"[📄 EXTRACTED] {len(extracted)} triples from {src}")
        total_triples += len(extracted)
        build_graph(extracted, source=src)
    
    print(f"[📊 GRAPH STATUS] Total triples: {total_triples}, Nodes: {G.number_of_nodes()}, Edges: {G.number_of_edges()}")

    print(f"[📊 GRAPH STATUS] Nodes: {G.number_of_nodes()}, Edges: {G.number_of_edges()}")
    
    if G.number_of_edges() == 0:
        print(f"[❌ GRAPH EMPTY] No edges found - using fallback")
        blank_img = visualize_graph(nx.DiGraph())
        info_df = pd.DataFrame([["No triples found (try relaxing strictness or enable fallback)"]], columns=["Info"])
        SEARCH_OPTIONS = []; GRAPH_CONTEXT_MEMORY="(no graph memory)"
        return blank_img, pd.DataFrame(), info_df, gr.update(choices=[], value=None), None, None, None, None, "No contract extraction (primary.py mode)"
    
    print(f"[🔄 CLUSTERING] Starting community detection...")

    # 5) Communities
    print(f"[🔄 CLUSTERING] Running community detection with resolution={LOUVAIN_RESOLUTION}...")
    print(f"[📊 PRE-CLUSTER] Graph has {G.number_of_nodes()} nodes, {G.number_of_edges()} edges before clustering")
    
    G, part = detect_communities(
        G, resolution=LOUVAIN_RESOLUTION, random_state=LOUVAIN_RANDOM_STATE,
        use_ontology_for_clustering=bool(use_onto_cluster_cb)
    )
    
    num_communities = len(set(part.values()))
    print(f"[✅ CLUSTERING] Found {num_communities} communities:")
    
    # Log community details
    community_sizes = {}
    for node, community in part.items():
        if community not in community_sizes:
            community_sizes[community] = []
        community_sizes[community].append(node)
    
    for comm_id, nodes in community_sizes.items():
        print(f"[🔗 COMMUNITY {comm_id}] {len(nodes)} nodes: {', '.join(nodes[:5])}{'...' if len(nodes) > 5 else ''}")

    # 6) Cluster summaries
    print(f"[📊 SUMMARIES] Generating summaries for {num_communities} clusters...")
    global partition
    partition = part
    CLUSTER_SUMMARIES.clear()
    
    print(f"[📊 SUMMARIES] Input text length: {len(RAW_COMBINED_TEXT)} chars")
    cluster_summaries = summarize_clusters(G, partition, RAW_COMBINED_TEXT)
    CLUSTER_SUMMARIES.update(cluster_summaries)
    
    print(f"[✅ SUMMARIES] Generated {len(CLUSTER_SUMMARIES)} cluster summaries:")
    for comm_id, summary in CLUSTER_SUMMARIES.items():
        summary_preview = summary[:100] + "..." if len(summary) > 100 else summary
        print(f"[📝 SUMMARY {comm_id}] {summary_preview}")
    
    print(f"[🎨 VISUALIZATION] Creating graph visualization...")
    graph_img = visualize_graph(G)
    print(f"[✅ VISUALIZATION] Graph visualization complete")

    summary_df = pd.DataFrame.from_dict(CLUSTER_SUMMARIES, orient='index').reset_index()
    summary_df = summary_df.rename(columns={"index": "Cluster Number"})
    cols = ["Cluster Number"] + [c for c in summary_df.columns if c != "Cluster Number"]
    summary_df = summary_df[cols]

    metrics_df = compute_metrics(G, partition)
    meta = pd.DataFrame([("Extracted triples", total_triples)], columns=["Metric", "Value"])
    metrics_df = pd.concat([meta, metrics_df], ignore_index=True)

    SEARCH_OPTIONS = _build_search_options(G, partition)
    search_update = gr.update(choices=SEARCH_OPTIONS, value=None)

    GRAPH_CONTEXT_MEMORY = build_graph_context_memory(G, partition, CLUSTER_SUMMARIES)

    # 7) Contract extraction AFTER graph building (CORRECT ORDER!)
    contract_status = "No contract extraction"
    print(f"[📊 CONTRACT EXTRACTION] Starting extraction AFTER graph building...")
    
    try:
        from contract_extractor import process_uploaded_documents_for_dashboard
        print(f"[📊 CONTRACT EXTRACTION] Extracting and merging with complete portfolio...")
        result = process_uploaded_documents_for_dashboard(states)
        
        if "error" in result:
            contract_status = f"Contract extraction failed: {result['error']}"
            print(f"[❌ CONTRACT EXTRACTION ERROR] {result['error']}")
        else:
            # Quick save of just the new contracts
            new_contracts_df = result["uploaded_contracts_df"]
            new_contracts_df.to_csv("new_contracts_only.csv", index=False)
            
            contract_status = f"✅ Extracted {result['new_contracts_count']} uploaded contracts, total portfolio: {result['total_contracts']}"
            print(f"[✅ CONTRACT EXTRACTION] {result['new_contracts_count']} contracts extracted from document")
            print(f"[✅ PORTFOLIO UPDATE] Total portfolio now has {result['total_contracts']} contracts")
            print(f"[💾 QUICK SAVE] Saved to new_contracts_only.csv")
    except Exception as e:
        contract_status = f"Contract extraction error: {str(e)}"
        print(f"[❌ CONTRACT EXTRACTION ERROR] {str(e)}")
        traceback.print_exc()

    print(f"[📤 EXPORT] Starting CFO analytics JSONL export...")
    # Auto-export CFO analytics for dashboard (with improved timeout handling)
    try:
        import sys
        sys.path.append('.')
        import config
        
        # Update global config for export function
        config.GRAPH_CONTEXT_MEMORY = GRAPH_CONTEXT_MEMORY
        config.RAW_COMBINED_TEXT = RAW_COMBINED_TEXT
        
        from export import export_cfo_jsonl_from_context
        
        print(f"[📤 INNER EXPORT] Calling export function...")
        print(f"[📊 GRAPH MEMORY] Context length: {len(GRAPH_CONTEXT_MEMORY)} chars")
        print(f"[📄 RAW TEXT] Text slice length: {len(RAW_COMBINED_TEXT[:3500])} chars")
        
        # Export JSONL with global config
        result = export_cfo_jsonl_from_context()
        
        if result:
            print(f"[✅ EXPORT] CFO JSONL exported to: {result}")
        else:
            print(f"[⚠️ EXPORT] CFO export returned no file")
            
    except Exception as e:
        print(f"[❌ EXPORT] CFO export failed: {e}")
        import traceback
        traceback.print_exc()

    # Save graph context memory to file for chatbot to use
    try:
        if GRAPH_CONTEXT_MEMORY and GRAPH_CONTEXT_MEMORY.strip() and GRAPH_CONTEXT_MEMORY != "(no graph memory)":
            context_file = "graph_context_memory.txt"
            with open(context_file, 'w', encoding='utf-8') as f:
                f.write(GRAPH_CONTEXT_MEMORY)
            print(f"[💾 SAVED] Graph context memory saved to {context_file} ({len(GRAPH_CONTEXT_MEMORY)} chars)")
            print(f"[✅ CHATBOT READY] Uploaded contracts are now available in chatbot context!")
        else:
            print(f"[⚠️ WARNING] Graph context memory is empty, not saving to file")
    except Exception as e:
        print(f"[⚠️ WARNING] Could not save graph context to file: {e}")
    
    print(f"[🎉 COMPLETE] Processing finished successfully!")
    # Extra outputs placeholders
    return graph_img, summary_df, metrics_df, search_update, None, None, None, None, contract_status

def process_corpus_and_build_graph(
    upload_paths,
    dir_path,
    type_choices,
    enable_vector=False,
    vector_topk=8,
    extract_temp=EXTRACT_TEMP,
    strict_prompt=STRICT_PROMPT,
    require_evidence=REQUIRE_EVIDENCE,
    enable_fallback=ENABLE_FALLBACK_COOCC,
    max_coocc_pairs=MAX_COOCC_PAIRS,
    chunk_size=CHUNK_MAX_CHARS,
    relation_policy=RELATION_POLICY,      # Off | Standard | Strict
    alias_norm=ENABLE_ALIAS_NORMALIZATION,
    skip_no_evidence=SKIP_EDGES_NO_EVID,
    louvain_resolution=LOUVAIN_RESOLUTION,
    louvain_random_state=LOUVAIN_RANDOM_STATE,
    use_onto_cluster_cb=False,
):
    global G, partition, ENABLE_VECTOR, VECTOR_TOPK, SEARCH_OPTIONS, CLUSTER_SUMMARIES, GRAPH_CONTEXT_MEMORY, RAW_COMBINED_TEXT
    global EXTRACT_TEMP, STRICT_PROMPT, REQUIRE_EVIDENCE, ENABLE_FALLBACK_COOCC, MAX_COOCC_PAIRS, CHUNK_MAX_CHARS
    global RELATION_POLICY, ENABLE_ALIAS_NORMALIZATION, SKIP_EDGES_NO_EVID, LOUVAIN_RESOLUTION, LOUVAIN_RANDOM_STATE

    ENABLE_VECTOR = bool(enable_vector)
    VECTOR_TOPK = int(vector_topk)
    EXTRACT_TEMP = float(extract_temp)
    STRICT_PROMPT = bool(strict_prompt)
    REQUIRE_EVIDENCE = bool(require_evidence)
    ENABLE_FALLBACK_COOCC = bool(enable_fallback)
    MAX_COOCC_PAIRS = int(max_coocc_pairs)
    CHUNK_MAX_CHARS = int(chunk_size)
    RELATION_POLICY = relation_policy
    ENABLE_ALIAS_NORMALIZATION = bool(alias_norm)
    SKIP_EDGES_NO_EVID = bool(skip_no_evidence)
    LOUVAIN_RESOLUTION = float(louvain_resolution)
    LOUVAIN_RANDOM_STATE = int(louvain_random_state)

    # 1) Load
    pairs = load_corpus(upload_paths, dir_path, type_choices)
    if not pairs:
        blank_img = visualize_graph(nx.DiGraph())
        info_df = pd.DataFrame([["No files were found/parsed. Check your directory path and type filters."]], columns=["Info"])
        SEARCH_OPTIONS = []; GRAPH_CONTEXT_MEMORY = "(no graph memory)"; RAW_COMBINED_TEXT = ""
        return blank_img, pd.DataFrame(), info_df, gr.update(choices=[], value=None), None, None, None, None

    RAW_COMBINED_TEXT = "\n\n".join([f"### {src}\n{txt}" for (src, txt) in pairs])

    # 2) Optional vector index
    if ENABLE_VECTOR:
        build_vector_index_from_docs(pairs)

    # 3) Extract & build graph
    G = nx.DiGraph()
    total_triples = 0
    for src, txt in pairs:
        extracted = extract_entities_relations(txt)
        total_triples += len(extracted)
        build_graph(extracted, source=src)

    if G.number_of_edges() == 0:
        blank_img = visualize_graph(nx.DiGraph())
        info_df = pd.DataFrame([["No triples found (try relaxing strictness or enable fallback)"]], columns=["Info"])
        SEARCH_OPTIONS = []; GRAPH_CONTEXT_MEMORY="(no graph memory)"
        return blank_img, pd.DataFrame(), info_df, gr.update(choices=[], value=None), None, None, None, None

    # 4) Communities
    G, part = detect_communities(
        G, resolution=LOUVAIN_RESOLUTION, random_state=LOUVAIN_RANDOM_STATE,
        use_ontology_for_clustering=bool(use_onto_cluster_cb)
    )

    # 5) Cluster summaries
    global partition
    partition = part
    CLUSTER_SUMMARIES.clear()
    CLUSTER_SUMMARIES.update(summarize_clusters(G, partition, RAW_COMBINED_TEXT))
    graph_img = visualize_graph(G)

    summary_df = pd.DataFrame.from_dict(CLUSTER_SUMMARIES, orient='index').reset_index()
    summary_df = summary_df.rename(columns={"index": "Cluster Number"})
    cols = ["Cluster Number"] + [c for c in summary_df.columns if c != "Cluster Number"]
    summary_df = summary_df[cols]

    metrics_df = compute_metrics(G, partition)
    meta = pd.DataFrame([("Extracted triples", total_triples)], columns=["Metric", "Value"])
    metrics_df = pd.concat([meta, metrics_df], ignore_index=True)

    SEARCH_OPTIONS = _build_search_options(G, partition)
    search_update = gr.update(choices=SEARCH_OPTIONS, value=None)

    GRAPH_CONTEXT_MEMORY = build_graph_context_memory(G, partition, CLUSTER_SUMMARIES)

    # Extra outputs placeholders
    return graph_img, summary_df, metrics_df, search_update, None, None, None, None

# ============ 7.1) CFO Export Helpers ============
def _compute_entity_table(G: nx.DiGraph) -> pd.DataFrame:
    nodes = list(G.nodes())
    communities = nx.get_node_attributes(G, 'community')
    in_deg = dict(G.in_degree())
    out_deg = dict(G.out_degree())
    rows = []
    for n in nodes:
        rows.append({
            "entity": n,
            "community": communities.get(n, None),
            "in_degree": int(in_deg.get(n, 0)),
            "out_degree": int(out_deg.get(n, 0)),
            "total_degree": int(in_deg.get(n, 0) + out_deg.get(n, 0)),
        })
    return pd.DataFrame(rows, columns=["entity","community","in_degree","out_degree","total_degree"])

def _compute_relationship_table(G: nx.DiGraph) -> pd.DataFrame:
    rows = []
    for u, v, d in G.edges(data=True):
        rows.append({
            "source_entity": u,
            "relation": d.get("label", ""),
            "target_entity": v,
            "evidence": d.get("evidence_text", ""),
            "source_file": d.get("source", ""),
            "source_community": G.nodes[u].get("community", None),
            "target_community": G.nodes[v].get("community", None),
        })
    cols = ["source_entity","relation","target_entity","evidence","source_file","source_community","target_community"]
    return pd.DataFrame(rows, columns=cols)

def _compute_cluster_table(G: nx.DiGraph, summaries: dict) -> pd.DataFrame:
    by_c = defaultdict(list)
    for n, c in nx.get_node_attributes(G, 'community').items():
        by_c[c].append(n)

    rows = []
    for cid, nodes in sorted(by_c.items(), key=lambda x: x[0]):
        sub = G.subgraph(nodes).copy()
        deg = dict(sub.degree())
        top_nodes = [n for n,_ in sorted(deg.items(), key=lambda x: x[1], reverse=True)[:8]]
        label = summaries.get(cid, {}).get("KPI/Risk Label", "")
        summ  = summaries.get(cid, {}).get("Cluster Summary", "")
        samples = []
        for u, v, d in itertools.islice(sub.edges(data=True), 8):
            samples.append(f"{u}[{d.get('label','')}]→{v}")
        rows.append({
            "cluster_id": cid,
            "kpi_risk_label": label,
            "cluster_summary": summ,
            "node_count": sub.number_of_nodes(),
            "edge_count": sub.number_of_edges(),
            "top_entities": "; ".join(top_nodes),
            "sample_relations": "; ".join(samples),
        })
    cols = ["cluster_id","kpi_risk_label","cluster_summary","node_count","edge_count","top_entities","sample_relations"]
    return pd.DataFrame(rows, columns=cols)

def export_cfo_csv_and_txt():
    if G.number_of_nodes() == 0:
        return None, None, None, None

    entities_df = _compute_entity_table(G)
    relationships_df = _compute_relationship_table(G)
    clusters_df = _compute_cluster_table(G, CLUSTER_SUMMARIES)

    ent_path = "entities.csv"
    rel_path = "relationships.csv"
    clu_path = "clusters.csv"
    entities_df.to_csv(ent_path, index=False, encoding="utf-8")
    relationships_df.to_csv(rel_path, index=False, encoding="utf-8")
    clusters_df.to_csv(clu_path, index=False, encoding="utf-8")

    txt_path = "cfo_dashboard_export.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("# ENTITIES (CSV)\n")
        entities_df.to_csv(f, index=False, encoding="utf-8")
        f.write("\n# RELATIONSHIPS (CSV)\n")
        relationships_df.to_csv(f, index=False, encoding="utf-8")
        f.write("\n# CLUSTERS (CSV)\n")
        clusters_df.to_csv(f, index=False, encoding="utf-8")

    return ent_path, rel_path, clu_path, txt_path

# ============ 7.2) CFO JSONL Export from Context Memory ============
CFO_DIMENSIONS_AND_QUESTIONS = [
    ("Financial Exposure & Obligations", "What is the total value of active contracts by business unit and geography?"),
    ("Financial Exposure & Obligations", "How much committed spend (future obligations) exists across all contracts?"),
    ("Financial Exposure & Obligations", "Which vendors/customers represent the top 10% of our contractual spend or revenue?"),
    ("Financial Exposure & Obligations", "How much deferred revenue or prepaid obligations are tied to active contracts?"),
    ("Financial Exposure & Obligations", "What is the breakdown of fixed vs variable pricing terms across contracts?"),
    ("Financial Exposure & Obligations", "What is the weighted average contract duration and renewal cycle?"),
    ("Financial Exposure & Obligations", "Are there material contracts with escalation clauses that will increase costs in the next 12–24 months?"),

    ("Revenue & Profitability", "How much revenue is recognized vs pending under IFRS-15/ASC-606 obligations?"),
    ("Revenue & Profitability", "Which contracts have discount or rebate structures impacting margins?"),
    ("Revenue & Profitability", "What percentage of revenue comes from contracts with auto-renewal vs manual renewal?"),
    ("Revenue & Profitability", "Are there contracts with unfavorable payment terms (e.g., >90 days receivable)?"),
    ("Revenue & Profitability", "What portion of contracts include revenue-sharing or contingent fee clauses?"),

    ("Risk & Compliance", "How many contracts are missing signatures, approvals, or documentation?"),
    ("Risk & Compliance", "Are there clauses exposing the company to penalties, liquidated damages, or termination risks?"),
    ("Risk & Compliance", "Which contracts are linked to compliance requirements (GDPR, SOX, ESG, etc.)?"),
    ("Risk & Compliance", "Do any contracts have currency/FX exposure clauses that may affect financials?"),
    ("Risk & Compliance", "How many contracts are nearing expiry without renegotiation (business continuity risk)?"),

    ("Operational & Vendor Management", "Which vendors/customers are most critical based on contract dependency (>20% supply/revenue)?"),
    ("Operational & Vendor Management", "How many contracts include SLAs with penalties for underperformance?"),
    ("Operational & Vendor Management", "Are there overlapping contracts with the same vendor that could be consolidated?"),
    ("Operational & Vendor Management", "Which contracts allow renegotiation, cost reductions, or early exit?"),
    ("Operational & Vendor Management", "Are there single-source dependency contracts without alternative suppliers/customers?"),

    ("Cash Flow & Working Capital", "What is the average payment term (days payable/receivable) across contracts?"),
    ("Cash Flow & Working Capital", "Which contracts drive the largest advance payments, deposits, or milestones?"),
    ("Cash Flow & Working Capital", "How many contracts have late payments or disputes impacting cash flow?"),

    ("Strategic Insights", "Which contracts align with strategic initiatives (cloud, ESG, digital transformation)?"),
    ("Strategic Insights", "How many contracts have revenue linked to performance metrics (KPIs, volume tiers)?"),
    ("Strategic Insights", "Which contracts create long-term liabilities (leases under IFRS-16/ASC-842, take-or-pay obligations)?"),
    ("Strategic Insights", "How many contracts have intellectual property (IP) transfer, license, or exclusivity clauses?"),
    ("Strategic Insights", "Which contracts contain non-compete, confidentiality, or restrictive covenants impacting growth?")
]

def _cfo_json_schema_instruction():
    return (
        "Return STRICT JSON (an array) with one object per question. "
        "Each object MUST follow this schema:\n"
        "{\n"
        '  "dimension": "string",\n'
        '  "question": "string",\n'
        '  "insight": "string",\n'
        '  "kpis": { "name": "value/string/number", ... },\n'
        '  "risks": ["string", ...],\n'
        '  "opportunities": ["string", ...],\n'
        '  "evidence": [{"snippet": "string", "source": "string"}],\n'
        '  "data_gaps": ["string", ...],\n'
        '  "confidence": 0.0\n'
        "}\n"
        "- confidence in [0,1].\n"
        "- Use only information grounded in the provided context. If unknown, leave fields minimal and add a data_gaps note.\n"
        "- Keep answers concise and CFO-ready.\n"
    )

def _build_cfo_prompt_from_context(context_blob: str) -> List[dict]:
    questions_block = "\n".join([f"- [{dim}] {q}" for dim, q in CFO_DIMENSIONS_AND_QUESTIONS])
    user_prompt = f"""
You are a Contract Analytics Expert assisting the CFO.

Context Memory (contract knowledge graph + evidence):
\"\"\"{context_blob}\"\"\" 

Task:
Answer the following CFO questions using ONLY the context above. If data is incomplete, state assumptions and data gaps explicitly.

Questions:
{questions_block}

Output:
{_cfo_json_schema_instruction()}
"""
    system_msg = (
        "You are an expert financial analyst. Be precise, conservative, and avoid speculation. "
        "Use short sentences and CFO-readable language. Ground answers strictly in context."
    )
    return [
        {"role": "system", "content": system_msg},
        {"role": "user", "content": user_prompt},
    ]

def _json_safe_load(raw: str):
    raw = (raw or "").strip()
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip(), flags=re.IGNORECASE | re.MULTILINE)
    try:
        return json.loads(raw)
    except Exception:
        return None

def _normalize_records(parsed):
    """
    Ensure we always return a list[dict] with required keys.
    """
    records = []
    if isinstance(parsed, dict):
        parsed = [parsed]
    if not isinstance(parsed, list):
        return records
    for item in parsed:
        if not isinstance(item, dict):
            continue
        obj = {
            "dimension": item.get("dimension", ""),
            "question": item.get("question", ""),
            "insight": item.get("insight", ""),
            "kpis": item.get("kpis", {}) or {},
            "risks": item.get("risks", []) or [],
            "opportunities": item.get("opportunities", []) or [],
            "evidence": item.get("evidence", []) or [],
            "data_gaps": item.get("data_gaps", []) or [],
            "confidence": float(item.get("confidence", 0.0) or 0.0),
        }
        records.append(obj)
    return records

def export_cfo_jsonl_from_context(jsonl_path: str = "cfo_contract_insights.jsonl"):
    """
    Uses GRAPH_CONTEXT_MEMORY (+ a small slice of RAW_COMBINED_TEXT) as input to the CFO prompt,
    calls the LLM once to answer all 30 questions, and writes JSONL for external dashboards.
    ENHANCED: Now also loads contract data for meaningful KPIs.
    """
    if not GRAPH_CONTEXT_MEMORY or GRAPH_CONTEXT_MEMORY.strip() == "(no graph memory)":
        return None  # No graph memory built yet.

    raw_slice = (RAW_COMBINED_TEXT or "")[:3500]
    context_blob = f"{GRAPH_CONTEXT_MEMORY}\n\nRAW_CONTEXT_SAMPLE:\n{raw_slice}"

    # CRITICAL FIX: Load contract data for meaningful KPIs
    contract_df = None
    try:
        # Always try to load the complete portfolio for full CFO analysis
        if os.path.exists("uploaded_contracts.csv"):
            contract_df = pd.read_csv("uploaded_contracts.csv")
            print(f"[Export] Using complete portfolio: {len(contract_df)} contracts for CFO analysis")
        else:
            print(f"[Export] No uploaded_contracts.csv found, CFO insights may be limited")
    except Exception as e:
        print(f"[Export] Could not load contract data: {e}")
    
    # Also try to include new contracts if available
    try:
        if os.path.exists("new_contracts_only.csv"):
            new_df = pd.read_csv("new_contracts_only.csv")
            if contract_df is not None:
                # Merge new contracts if not already included
                existing_ids = set(contract_df['contract_id'].tolist()) if 'contract_id' in contract_df.columns else set()
                new_contracts = new_df[~new_df['contract_id'].isin(existing_ids)]
                if not new_contracts.empty:
                    contract_df = pd.concat([contract_df, new_contracts], ignore_index=True)
                    print(f"[Export] Added {len(new_contracts)} additional new contracts")
            else:
                contract_df = new_df
                print(f"[Export] Using only new contracts data: {len(contract_df)} contracts")
    except Exception as e:
        print(f"[Export] Could not merge new contracts: {e}")

    # Use the enhanced CFO analytics function that includes contract data
    try:
        import sys
        sys.path.append('.')
        from cfo_analytics import generate_cfo_insights_from_context
        
        records = generate_cfo_insights_from_context(context_blob, contract_df)
        
        if not records:
            # Fallback to basic records if generation fails
            records = []
            for dim, q in CFO_DIMENSIONS_AND_QUESTIONS:
                records.append({
                    "dimension": dim,
                    "question": q,
                    "insight": "",
                    "kpis": {},
                    "risks": [],
                    "opportunities": [],
                    "evidence": [],
                    "data_gaps": ["No extractable data from context memory."],
                    "confidence": 0.0,
                })
    except Exception as e:
        print(f"[Export] Error generating CFO insights: {e}")
        # Fallback to basic prompt-based generation
        messages = _build_cfo_prompt_from_context(context_blob)
        raw = _chat_strict(messages, temperature=0.1, timeout=60, model="gpt-4o")
        parsed = _json_safe_load(raw)
        records = _normalize_records(parsed)

        if not records:
            for dim, q in CFO_DIMENSIONS_AND_QUESTIONS:
                records.append({
                    "dimension": dim,
                    "question": q,
                    "insight": "",
                    "kpis": {},
                    "risks": [],
                    "opportunities": [],
                    "evidence": [],
                    "data_gaps": ["No extractable data from context memory."],
                    "confidence": 0.0,
                })

    run_meta = {
        "generated_at": pd.Timestamp.utcnow().isoformat(),
        "source": "GraphRAG Studio",
        "model": "gpt-4o"
    }

    try:
        with open(jsonl_path, "w", encoding="utf-8") as f:
            for r in records:
                obj = {**r, **{"_meta": run_meta}}
                f.write(json.dumps(obj, ensure_ascii=False) + "\n")
        return jsonl_path
    except Exception as e:
        print("JSONL write error:", e)
        return None

# ============ 8) Chat helpers ============
def _cluster_id_from_text(txt: str):
    m = re.search(r"\bcluster\s*([0-9]+)\b", txt, re.I)
    return int(m.group(1)) if m else None

def _summarize_cluster_by_id(cid: int):
    if not partition:
        return "No clusters yet. Run the pipeline first."
    nodes = [n for n, c in partition.items() if c == cid]
    if not nodes:
        return f"No such cluster: {cid}"

    s = CLUSTER_SUMMARIES.get(cid, {})
    label = s.get("KPI/Risk Label", "")
    summary = s.get("Cluster Summary", "")
    if not (label or summary):
        label, summary = _heuristic_cluster_label_and_summary(G, cid, nodes, "")

    sub = G.subgraph(nodes).copy()
    deg = dict(sub.degree())
    top_nodes = [n for n, _d in sorted(deg.items(), key=lambda x: x[1], reverse=True)[:8]]

    samples = []
    for u, v, d in itertools.islice(sub.edges(data=True), 6):
        lbl = d.get("label", "")
        ev = (d.get("evidence_text", "") or "").strip()
        if len(ev) > 160: ev = ev[:160].rstrip() + "…"
        samples.append(f"- {u} —[{lbl}]→ {v}; Ev: {ev}")

    lines = [f"**Cluster {cid}**",
             f"- **Label:** {label or '—'}",
             f"- **Summary:** {summary or '—'}"]
    if top_nodes:
        lines.append(f"- **Key entities:** {', '.join(top_nodes)}")
    if samples:
        lines.append("**Sample relations:**")
        lines.extend(samples)
    return "\n".join(lines)

def _describe_entity_freeform(text: str):
    m = re.search(
        r"^(?:describe|show|detail|explain)\s+(?:entities?\s+(?:for|of|from)\s+)?(.+)$",
        text.strip(), re.I
    )
    if m:
        target = m.group(1).strip()
        mcl = re.match(r"^cluster\s*([0-9]+)$", target, re.I)
        if mcl:
            cid = int(mcl.group(1))
            return _summarize_cluster_by_id(cid)
        return describe_selection(target)

    m2 = re.search(
        r"^(?:what\s+are\s+)?entities?\s+(?:for|of|from)\s+(.+)$",
        text.strip(), re.I
    )
    if m2:
        target = m2.group(1).strip()
        return describe_selection(target)

    return None

# ============ 8.1) LLM-grounded Q&A ============
def _tokenize_question(q):
    toks = re.findall(r"[a-z0-9][a-z0-9\-_/\.]{2,}", (q or "").lower())
    seen, out = set(), []
    for t in toks:
        if t not in seen:
            seen.add(t); out.append(t)
    return out

def _edge_score_for_question(u, v, d, q_tokens):
    rel = (d.get("label") or "").lower()
    ev  = (d.get("evidence_text") or "").lower()
    u_l = str(u).lower(); v_l = str(v).lower()
    def has(tok, text): return (tok in text)
    name_hits = sum(has(t, u_l) or has(t, v_l) for t in q_tokens)
    rel_hits  = sum(has(t, rel) for t in q_tokens)
    ev_hits   = sum(has(t, ev)  for t in q_tokens)
    score = (EDGE_NAME_WEIGHT * name_hits) + (EDGE_REL_WEIGHT * rel_hits) + (EDGE_EVID_WEIGHT * ev_hits)
    q_join = " ".join(q_tokens)
    if "recogn" in q_join and ("recogn" in rel or "revenue" in ev): score += 2
    if ("quarter" in q_join or "q1" in q_join or "q2" in q_join or "q3" in q_join or "q4" in q_join) and any(x in ev for x in ["q1","q2","q3","q4","quarter"]): score += 1
    return score

def _select_relevant_edges(question, G, top_k=MAX_EDGES_IN_PROMPT):
    q_tokens = _tokenize_question(question)
    if not q_tokens:
        deg = dict(G.degree())
        edges = sorted(G.edges(data=True), key=lambda e: -(deg.get(e[0],0)+deg.get(e[1],0)))
        return edges[:top_k]
    scored = []
    for u, v, d in G.edges(data=True):
        s = _edge_score_for_question(u, v, d, q_tokens)
        if s > 0: scored.append((s, u, v, d))
    if not scored:
        deg = dict(G.degree())
        edges = sorted(G.edges(data=True), key=lambda e: -(deg.get(e[0],0)+deg.get(e[1],0)))
        return edges[:min(12, top_k)]
    scored.sort(key=lambda x: -x[0])
    return [(u, v, d) for (s, u, v, d) in scored[:top_k]]

def _cluster_card(cid, G, partition, summaries, max_nodes=8):
    nodes = [n for n, c in partition.items() if c == cid]
    sub = G.subgraph(nodes).copy()
    deg = dict(sub.degree())
    top_nodes = [n for n,_ in sorted(deg.items(), key=lambda x: x[1], reverse=True)[:max_nodes]]
    s = summaries.get(cid, {})
    label = s.get("KPI/Risk Label", "") or ""
    summ  = s.get("Cluster Summary", "") or ""
    card  = [f"[Cluster {cid}] Label: {label}"]
    if summ: card.append(f"[Cluster {cid}] Summary: {summ}")
    if top_nodes: card.append(f"[Cluster {cid}] Key entities: {', '.join(top_nodes)}")
    return "\n".join(card)

def _build_llm_context(question, G, partition, summaries, enable_vector, vector_topk):
    edges = _select_relevant_edges(question, G, top_k=MAX_EDGES_IN_PROMPT)
    touched_cids = set()
    for u, v, _d in edges:
        if u in partition: touched_cids.add(partition[u])
        if v in partition: touched_cids.add(partition[v])
    cluster_blocks = []
    for cid in sorted(touched_cids):
        cluster_blocks.append(_cluster_card(cid, G, partition, summaries))
    edge_lines = []
    for u, v, d in edges:
        lbl = d.get("label","")
        ev  = (d.get("evidence_text","") or "").strip()
        if len(ev) > MAX_EVIDENCE_CHARS: ev = ev[:MAX_EVIDENCE_CHARS].rstrip() + "…"
        edge_lines.append(f"- {u} —[{lbl}]→ {v}; Ev: {ev}")
    vec_text = ""
    if enable_vector:
        try:
            bits = vector_snippets(question, top_k=vector_topk)
            if bits:
                vec_text = "\n\nAdditional snippets:\n" + "\n".join(f"- {b}" for b in bits[:6])
        except Exception:
            pass
    ctx = []
    if cluster_blocks: ctx.append("Cluster context:\n" + "\n\n".join(cluster_blocks))
    if edge_lines: ctx.append("Relevant relations:\n" + "\n".join(edge_lines))
    if vec_text: ctx.append(vec_text)
    return "\n\n".join(ctx), len(edge_lines)

def graph_chat(message, history=[]):
    global G, partition, ENABLE_VECTOR, VECTOR_TOPK, CLUSTER_SUMMARIES, GRAPH_CONTEXT_MEMORY
    if G.number_of_nodes() == 0:
        return "It looks like we haven't built a knowledge graph yet. Please add files and click Run."

    text = (message or "").strip()
    m = text.lower()

    if m.startswith("trace "):
        try:
            inner = m[len("trace "):]
            src, tgt = [x.strip() for x in inner.split(" to ", 1)]
            path = nx.shortest_path(G, source=src, target=tgt)
            return f"Found a path! Here it is: **{' → '.join(path)}**"
        except Exception:
            return "Sorry, I couldn't find a direct path between those two entities."

    if "connected to" in m:
        try:
            entity = m.split("connected to", 1)[-1].strip()
            if entity not in G:
                return f"I couldn't find an entity named '{entity}' in the graph."
            neighbors = list(G.successors(entity))
            predecessors = list(G.predecessors(entity))
            response_lines = [f"**{entity}** connections:"]
            if neighbors:
                response_lines.append(f"- Outgoing: {', '.join(neighbors)}")
            if predecessors:
                response_lines.append(f"- Incoming: {', '.join(predecessors)}")
            if not neighbors and not predecessors:
                return f"**{entity}** exists in the graph but has no direct connections."
            return "\n".join(response_lines)
        except Exception:
            return "Use: `connected to <Entity Name>`."

    norm = re.sub(r"\bcluster(\d+)\b", r"cluster \1", m)
    if re.search(r"\b(summarize|summary|summarise|describe)\s+cluster\s*\d+\b", norm):
        cid = _cluster_id_from_text(norm)
        if cid is not None:
            return _summarize_cluster_by_id(cid)

    if ("cluster summary" in m) or ("summarize clusters" in m) or ("summary of clusters" in m):
        if not CLUSTER_SUMMARIES: return "No cluster summaries are available yet."
        lines = ["Cluster summaries:"]
        for cid in sorted(CLUSTER_SUMMARIES.keys()):
            s = CLUSTER_SUMMARIES[cid]
            label = s.get("KPI/Risk Label", "N/A")
            summ  = s.get("Cluster Summary", "No summary.")
            lines.append(f"\n**Cluster {cid}: {label}**\n{summ}")
        return "\n".join(lines)

    desc = _describe_entity_freeform(text)
    if desc:
        return desc

    llm_context, edge_count = _build_llm_context(
        text, G, partition, CLUSTER_SUMMARIES, ENABLE_VECTOR, VECTOR_TOPK
    )

    if edge_count == 0 and not (ENABLE_VECTOR and "Additional snippets" in llm_context):
        return "I couldn't find specific relations for that. Try asking about a main entity or cluster."

    system_msg = (
        "You are a helpful financial analyst assistant named 'GraphAnalyst'. "
        "Base answers strictly on the provided context derived from the user's knowledge graph."
    )

    full_context = f"Context:\n{GRAPH_CONTEXT_MEMORY}\n\n{llm_context}"
    user_msg = f"{full_context}\n\nQuestion: {text}\n\nGive a concise, direct answer grounded in the context."

    messages = [
        {"role": "system", "content": system_msg},
        {"role": "user", "content": user_msg}
    ]

    out = _chat_strict(messages, temperature=0.2, timeout=45, model="gpt-4o")
    if not out:
        facts = [ln for ln in llm_context.splitlines() if ln.startswith("- ")]
        if facts:
            return "Relevant facts:\n" + "\n".join(facts[:5])
        return "Sorry, I couldn't formulate an answer from the context."
    return out

# ============ 8.5) Search / Autocomplete describe ============
def describe_selection(selection):
    if not selection:
        return "Select an entity or a cluster."
    sel = str(selection).strip()
    if sel.lower().startswith("cluster "):
        try:
            cid = int(sel.split(" ", 1)[1])
        except Exception:
            return f"Could not parse cluster id from '{sel}'."
        nodes = [n for n, c in partition.items() if c == cid]
        if not nodes: return f"No such cluster: {cid}"
        sub = G.subgraph(nodes).copy(); edges = sub.number_of_edges()
        degrees = dict(sub.degree())
        top_deg = sorted(degrees.items(), key=lambda x: x[1], reverse=True)[:10]
        md = []
        md.append(f"### Cluster {cid}")
        md.append(f"- Nodes: **{len(nodes)}**, Edges (internal): **{edges}**")
        md.append("- Top nodes by degree: " + ", ".join([f"**{n}** ({d})" for n, d in top_deg]) if top_deg else "-")
        samples = []
        for u, v, d in itertools.islice(sub.edges(data=True), 10):
            lbl = d.get("label", ""); ev = d.get("evidence_text", "")
            samples.append(f"- **{u}** --[{lbl}]--> **{v}** \n  Evidence: {ev}")
        if samples:
            md.append("#### Sample relations"); md.extend(samples)
        return "\n".join(md)

    node = None
    if sel in G:
        node = sel
    else:
        cand = [n for n in G.nodes() if sel.lower() in str(n).lower()]
        if cand: node = cand[0]
    if node is None:
        return f"No entity found for '{sel}'."

    cid = G.nodes[node].get("community", "(n/a)")
    in_neighbors = list(G.predecessors(node)); out_neighbors = list(G.successors(node))
    deg = G.degree(node)
    md = []
    md.append(f"### Entity: **{node}**")
    md.append(f"- Cluster: **{cid}**")
    md.append(f"- Degree: **{deg}** (in: {len(in_neighbors)}, out: {len(out_neighbors)})")
    if out_neighbors:
        md.append("#### Outgoing connections")
        for v in out_neighbors[:15]:
            d = G.get_edge_data(node, v, default={})
            lbl = d.get("label", ""); ev = d.get("evidence_text", "")
            md.append(f"- → **{v}** [{lbl}]  \n  Evidence: {ev}")
    if in_neighbors:
        md.append("#### Incoming connections")
        for u in in_neighbors[:15]:
            d = G.get_edge_data(u, node, default={})
            lbl = d.get("label", ""); ev = d.get("evidence_text", "")
            md.append(f"- ← **{u}** [{lbl}]  \n  Evidence: {ev}")
    return "\n".join(md)

# ============ 9) RDF / Ontology helpers ============
DEFAULT_BASE_IRI = "http://example.org/kg/"
DEFAULT_ONTO_IRI = "http://example.org/onto/"
DEFAULT_CONTEXT_IRI = "http://example.org/context/"
DEFAULT_ONTOLOGY_FILENAME = "business_ontology.ttl"
DEFAULT_RDF_TTL_FILENAME = "graph_export.ttl"
DEFAULT_JSONLD_FILENAME = "graph_export.jsonld"

REL_TO_PROP = {
    "audited_by": "auditedBy",
    "governed_by": "governedBy",
    "recognised_under": "recognisedUnder",
    "has_obligation": "hasObligation",
    "involves_counterparty": "involvesCounterparty",
    "maps_to_standard": "mapsToStandard",
    "linked_to": "linkedTo",
    "evidenced_by": "evidencedBy",
    "co_occurs_in_sentence": "coOccursInSentence",
}

def _slug(name: str) -> str:
    return quote(re.sub(r"[^A-Za-z0-9_\-\.]+", "_", name.strip()), safe="_-.")

def _node_iri(node: str, base_iri: str) -> str:
    return base_iri.rstrip("/") + "/entity/" + _slug(node)

def build_business_ontology(onto_iri: str = DEFAULT_ONTO_IRI, base_iri: str = DEFAULT_BASE_IRI):
    if not RDFLIB_AVAILABLE:
        raise RuntimeError("rdflib not installed; cannot build ontology. pip install rdflib")
    g = RDFGraph()
    ONTO = Namespace(onto_iri if onto_iri.endswith(('#','/')) else onto_iri + "#")
    EX   = Namespace(base_iri if base_iri.endswith(('#','/')) else base_iri + "#")
    g.bind("owl", OWL); g.bind("rdfs", RDFS); g.bind("rdf", RDF)
    g.bind("skos", SKOS); g.bind("dct", DCTERMS)
    g.bind("onto", ONTO); g.bind("ex", EX)

    onto_uri = URIRef(onto_iri.rstrip("#/"))
    g.add((onto_uri, RDF.type, OWL.Ontology))
    g.add((onto_uri, RDFS.label, Literal("Business Ontology for GraphRAG", lang="en")))

    Entity = ONTO.Entity
    Standard = ONTO.Standard
    g.add((Entity, RDF.type, OWL.Class)); g.add((Entity, RDFS.label, Literal("Entity", lang="en")))
    g.add((Standard, RDF.type, OWL.Class)); g.add((Standard, RDFS.label, Literal("Standard", lang="en")))

    def add_obj_prop(local, label, domain=Entity, range_=Entity):
        p = ONTO[local]
        g.add((p, RDF.type, OWL.ObjectProperty))
        g.add((p, RDFS.label, Literal(label, lang="en")))
        g.add((p, RDFS.domain, domain)); g.add((p, RDFS.range, range_))
        return p

    add_obj_prop("auditedBy", "audited by")
    add_obj_prop("governedBy", "governed by")
    add_obj_prop("recognisedUnder", "recognised under", domain=Entity, range_=Standard)
    add_obj_prop("hasObligation", "has obligation")
    add_obj_prop("involvesCounterparty", "involves counterparty")
    add_obj_prop("mapsToStandard", "maps to standard", domain=Entity, range_=Standard)
    add_obj_prop("linkedTo", "linked to")
    add_obj_prop("evidencedBy", "evidenced by")
    add_obj_prop("coOccursInSentence", "co-occurs in sentence")

    EX_ASC606 = EX.ASC_606; EX_ASC842 = EX.ASC_842
    g.add((EX_ASC606, RDF.type, Standard)); g.add((EX_ASC606, SKOS.prefLabel, Literal("ASC 606")))
    g.add((EX_ASC842, RDF.type, Standard)); g.add((EX_ASC842, SKOS.prefLabel, Literal("ASC 842")))

    ttl_path = DEFAULT_ONTOLOGY_FILENAME
    g.serialize(destination=ttl_path, format="turtle")
    return ttl_path

def export_graph_as_rdf(base_iri=DEFAULT_BASE_IRI, onto_iri=DEFAULT_ONTO_IRI, context_iri=DEFAULT_CONTEXT_IRI):
    """Export graph as enhanced RDF with CFO analytics extensions including KPIs"""
    if not RDFLIB_AVAILABLE:
        raise RuntimeError("rdflib not installed; cannot export RDF. pip install rdflib")
    g = RDFGraph()
    ONTO = Namespace(onto_iri if onto_iri.endswith(('#','/')) else onto_iri + "#")
    EX   = Namespace(base_iri if base_iri.endswith(('#','/')) else base_iri + "#")
    g.bind("onto", ONTO); g.bind("ex", EX); g.bind("dct", DCTERMS); g.bind("skos", SKOS)
    g.bind("cfo", Namespace("http://example.org/cfo#"))

    # Define CFO-specific classes
    CFO_NS = Namespace("http://example.org/cfo#")
    g.add((ONTO.Entity, RDF.type, OWL.Class))
    g.add((ONTO.Standard, RDF.type, OWL.Class))
    g.add((CFO_NS.CFOInsight, RDF.type, OWL.Class))
    g.add((CFO_NS.KPI, RDF.type, OWL.Class))
    g.add((CFO_NS.Risk, RDF.type, OWL.Class))
    g.add((CFO_NS.Opportunity, RDF.type, OWL.Class))

    for node, attrs in G.nodes(data=True):
        uri = URIRef(_node_iri(node, base_iri))
        g.add((uri, RDF.type, ONTO.Entity))
        g.add((uri, SKOS.prefLabel, Literal(str(node))))
        cid = attrs.get("community", None)
        if cid is not None:
            g.add((uri, DCTERMS.subject, Literal(f"cluster:{cid}")))

    prop_cache = {}
    def prop_for(rel_norm: str):
        local = REL_TO_PROP.get(rel_norm, rel_norm)
        if local not in prop_cache:
            prop_cache[local] = ONTO[local]
            g.add((prop_cache[local], RDF.type, OWL.ObjectProperty))
            g.add((prop_cache[local], RDFS.label, Literal(local)))
        return prop_cache[local]

    for u, v, d in G.edges(data=True):
        rel = (d.get("label") or "").lower()
        p = prop_for(rel)
        s = URIRef(_node_iri(u, base_iri)); o = URIRef(_node_iri(v, base_iri))
        g.add((s, p, o))
        ev = d.get("evidence_text", ""); src = d.get("source", "")
        if ev:
            b = BNode()
            g.add((s, ONTO.evidencedBy, b))
            g.add((b, RDF.type, RDFS.Resource))
            g.add((b, DCTERMS.description, Literal(ev)))
            if src:
                g.add((b, DCTERMS.source, Literal(src)))

    # CRITICAL FIX: Add CFO insights and KPIs to the RDF graph
    try:
        # Load CFO insights from JSONL file
        cfo_insights = []
        if os.path.exists("cfo_contract_insights.jsonl"):
            with open("cfo_contract_insights.jsonl", "r", encoding="utf-8") as f:
                for line in f:
                    if line.strip():
                        try:
                            insight = json.loads(line.strip())
                            cfo_insights.append(insight)
                        except json.JSONDecodeError:
                            continue
        
        # Add CFO insights to RDF graph
        for i, insight in enumerate(cfo_insights):
            insight_uri = URIRef(f"{base_iri}cfo_insight_{i}")
            g.add((insight_uri, RDF.type, CFO_NS.CFOInsight))
            g.add((insight_uri, DCTERMS.title, Literal(insight.get("dimension", "Unknown"))))
            g.add((insight_uri, DCTERMS.description, Literal(insight.get("question", ""))))
            g.add((insight_uri, CFO_NS.insight, Literal(insight.get("insight", ""))))
            g.add((insight_uri, CFO_NS.confidence, Literal(insight.get("confidence", 0.0))))
            
            # Add KPIs
            kpis = insight.get("kpis", {})
            for kpi_name, kpi_value in kpis.items():
                # URL encode KPI names to handle spaces and special characters
                import urllib.parse
                safe_kpi_name = urllib.parse.quote(kpi_name, safe='')
                kpi_uri = URIRef(f"{base_iri}kpi_{i}_{safe_kpi_name}")
                g.add((kpi_uri, RDF.type, CFO_NS.KPI))
                g.add((kpi_uri, DCTERMS.title, Literal(kpi_name)))
                g.add((kpi_uri, CFO_NS.value, Literal(str(kpi_value))))
                g.add((insight_uri, CFO_NS.hasKPI, kpi_uri))
            
            # Add risks
            risks = insight.get("risks", [])
            for j, risk in enumerate(risks):
                risk_uri = URIRef(f"{base_iri}risk_{i}_{j}")
                g.add((risk_uri, RDF.type, CFO_NS.Risk))
                g.add((risk_uri, DCTERMS.description, Literal(risk)))
                g.add((insight_uri, CFO_NS.hasRisk, risk_uri))
            
            # Add opportunities
            opportunities = insight.get("opportunities", [])
            for j, opp in enumerate(opportunities):
                opp_uri = URIRef(f"{base_iri}opportunity_{i}_{j}")
                g.add((opp_uri, RDF.type, CFO_NS.Opportunity))
                g.add((opp_uri, DCTERMS.description, Literal(opp)))
                g.add((insight_uri, CFO_NS.hasOpportunity, opp_uri))
            
            # Add evidence
            evidence = insight.get("evidence", [])
            for j, ev in enumerate(evidence):
                ev_uri = URIRef(f"{base_iri}evidence_{i}_{j}")
                g.add((ev_uri, RDF.type, RDFS.Resource))
                g.add((ev_uri, DCTERMS.description, Literal(ev.get("snippet", ""))))
                g.add((ev_uri, DCTERMS.source, Literal(ev.get("source", ""))))
                g.add((insight_uri, CFO_NS.evidencedBy, ev_uri))
        
        print(f"[Export] Added {len(cfo_insights)} CFO insights with KPIs to RDF graph")
        
    except Exception as e:
        print(f"[Export] Warning: Could not add CFO insights to RDF: {e}")

    ttl_path = DEFAULT_RDF_TTL_FILENAME
    jsonld_path = DEFAULT_JSONLD_FILENAME
    g.serialize(destination=ttl_path, format="turtle")
    try:
        g.serialize(destination=jsonld_path, format="json-ld")
        print(f"[Export] Successfully exported enhanced JSON-LD with CFO insights and KPIs")
    except Exception as e:
        print(f"[Export] Error serializing JSON-LD: {e}")
        jsonld_path = None
    return ttl_path, jsonld_path

# --- Bridge for ChatInterface(type="messages") ---
def _chat_bridge(messages, _history=None):
    last_user = ""
    if isinstance(messages, list):
        for m in reversed(messages):
            if isinstance(m, dict) and m.get("role") == "user":
                last_user = m.get("content", "")
                break
            if isinstance(m, (list, tuple)) and len(m) >= 2 and m[0] == "user":
                last_user = m[1]
                break
    elif isinstance(messages, str):
        last_user = messages
    return graph_chat(last_user, [])

# ============ 10) UI ============
with gr.Blocks() as demo:
    gr.Markdown("## Contract GraphRAG Studio")

    with gr.Row():
        file_input = gr.Files(
            label="Upload documents (.pdf, .docx, .txt)",
            file_types=[".pdf", ".docx", ".txt"],
        )
        dir_input = gr.Textbox(
            label="OR: Directory path (will recurse)",
            placeholder="e.g., C:/Users/Me/Documents/contracts"
        )

    with gr.Row():
        type_selector = gr.CheckboxGroup(
            choices=["PDF","Word","Text"],
            value=["PDF","Word","Text"],
            label="Include file types"
        )

    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### Vector Search")
            enable_vector_cb = gr.Checkbox(value=False, label="Enable Vector Search (LlamaIndex)")
            topk_slider = gr.Slider(minimum=2, maximum=16, value=8, step=1, label="Vector top-k")

            gr.Markdown("### RDF / Ontology (optional)")
            base_iri_tb = gr.Textbox(value=DEFAULT_BASE_IRI, label="Base IRI (entities)")
            onto_iri_tb = gr.Textbox(value=DEFAULT_ONTO_IRI, label="Ontology IRI")
            use_onto_cluster_cb = gr.Checkbox(value=False, label="Use ontology-weighted clustering")
            build_onto_btn = gr.Button("Build Ontology (TTL)")
            export_rdf_btn = gr.Button("Export Graph (TTL + JSON-LD)")
            onto_file_out = gr.File(label="Ontology (TTL)", interactive=False)
            rdf_file_out = gr.File(label="Graph TTL", interactive=False)
            jsonld_file_out = gr.File(label="Graph JSON-LD", interactive=False)

            badge_rdf = "✅ rdflib loaded" if RDFLIB_AVAILABLE else "⚠️ rdflib not installed"
            badge_pdf = "✅ PyPDF2" if PDF_AVAILABLE else "⚠️ PyPDF2 missing (PDFs skipped)"
            badge_docx = "✅ python-docx" if DOCX_AVAILABLE else "⚠️ python-docx missing (.docx skipped)"
            gr.Markdown(f"**Parsers:** {badge_pdf} • {badge_docx}  \n**RDF:** {badge_rdf}")

        with gr.Column(scale=2):
            gr.Markdown("### Graph Tuning")
            with gr.Row():
                extract_temp = gr.Slider(0.0, 1.0, value=EXTRACT_TEMP, step=0.05, label="Extraction temperature")
                strict_prompt = gr.Checkbox(value=STRICT_PROMPT, label="Strict prompt (reject weak edges)")
                require_evidence = gr.Checkbox(value=REQUIRE_EVIDENCE, label="Require evidence on edges")
            with gr.Row():
                enable_fallback = gr.Checkbox(value=ENABLE_FALLBACK_COOCC, label="Enable fallback co-occurrence")
                max_coocc_pairs = gr.Slider(0, 10, value=MAX_COOCC_PAIRS, step=1, label="Max co-occurrence pairs/sentence")
            chunk_size = gr.Slider(2000, 5000, value=CHUNK_MAX_CHARS, step=100, label="Chunk size (chars)")
            with gr.Row():
                relation_policy = gr.Dropdown(choices=["Off","Standard","Strict"], value=RELATION_POLICY, label="Relation whitelist")
                alias_norm = gr.Checkbox(value=ENABLE_ALIAS_NORMALIZATION, label="Alias normalization")
                skip_no_evidence = gr.Checkbox(value=SKIP_EDGES_NO_EVID, label="Skip edges without evidence")
            with gr.Row():
                louvain_resolution = gr.Slider(0.2, 2.0, value=LOUVAIN_RESOLUTION, step=0.05, label="Louvain resolution")
                louvain_random_state = gr.Slider(1, 999, value=LOUVAIN_RANDOM_STATE, step=1, label="Random state")

    run_button = gr.Button("Run: Build Graph, Clusters & Metrics")

    with gr.Row():
        graph_output = gr.Image(label="🕸️ Knowledge Graph")
        summary_output = gr.Dataframe(label="🧾 Cluster Summaries", interactive=False)
        metrics_output = gr.Dataframe(label="📈 Graph Metrics", interactive=False)
    
    # Contract extraction status
    contract_status_output = gr.Textbox(label="📋 Contract Extraction Status", interactive=False)

    gr.Markdown("### 🔎 Search Entity or Cluster")
    with gr.Row():
        search_dd = gr.Dropdown(label="Type to filter…", choices=[], value=None, allow_custom_value=True, interactive=True)
        search_btn = gr.Button("Show Details")
    search_md = gr.Markdown()

    gr.Markdown("### 📤 Export for CFO Dashboard")
    export_cfo_btn = gr.Button("Export CFO CSV/TXT")
    entities_csv_out = gr.File(label="entities.csv", interactive=False)
    relationships_csv_out = gr.File(label="relationships.csv", interactive=False)
    clusters_csv_out = gr.File(label="clusters.csv", interactive=False)
    cfo_txt_out = gr.File(label="cfo_dashboard_export.txt", interactive=False)

    # --- NEW: JSONL export button and wiring ---
    gr.Markdown("### 🧠 CFO Insights (JSONL for External Streamlit Dashboard)")
    export_cfo_jsonl_btn = gr.Button("Export CFO JSONL (for Dashboard)")
    cfo_jsonl_out = gr.File(label="cfo_contract_insights.jsonl", interactive=False)

    # Wire main pipeline
    run_button.click(
        fn=process_corpus_with_contract_extraction,
        inputs=[
            file_input,           # Files
            dir_input,            # Directory path
            type_selector,        # Types
            enable_vector_cb, topk_slider,
            extract_temp, strict_prompt, require_evidence, enable_fallback, max_coocc_pairs,
            chunk_size, relation_policy, alias_norm, skip_no_evidence,
            louvain_resolution, louvain_random_state,
            use_onto_cluster_cb
        ],
        outputs=[graph_output, summary_output, metrics_output, search_dd, rdf_file_out, jsonld_file_out, entities_csv_out, cfo_txt_out, contract_status_output]
    )

    search_btn.click(fn=describe_selection, inputs=[search_dd], outputs=[search_md])

    # RDF buttons
    def _build_onto_action(onto_iri, base_iri):
        if not RDFLIB_AVAILABLE:
            return None, gr.update(value=False)
        try:
            path = build_business_ontology(onto_iri=onto_iri, base_iri=base_iri)
            return path, gr.update(value=True)
        except Exception as e:
            print("Ontology build error:", e)
            return None, gr.update(value=False)

    def _export_rdf_action(base_iri, onto_iri):
        if not RDFLIB_AVAILABLE:
            return None, None
        try:
            ttl, jsonld = export_graph_as_rdf(base_iri=base_iri, onto_iri=onto_iri)
            return ttl, jsonld
        except Exception as e:
            print("RDF export error:", e)
            return None, None

    build_onto_btn.click(
        fn=_build_onto_action,
        inputs=[onto_iri_tb, base_iri_tb],
        outputs=[onto_file_out, use_onto_cluster_cb]
    )
    export_rdf_btn.click(
        fn=_export_rdf_action,
        inputs=[base_iri_tb, onto_iri_tb],
        outputs=[rdf_file_out, jsonld_file_out]
    )

    # CFO export
    export_cfo_btn.click(
        fn=export_cfo_csv_and_txt,
        inputs=[],
        outputs=[entities_csv_out, relationships_csv_out, clusters_csv_out, cfo_txt_out]
    )

    # CFO JSONL export wiring
    def _export_cfo_jsonl_action():
        path = export_cfo_jsonl_from_context(jsonl_path="cfo_contract_insights.jsonl")
        return path

    export_cfo_jsonl_btn.click(
        fn=_export_cfo_jsonl_action,
        inputs=[],
        outputs=[cfo_jsonl_out]
    )

    gr.Markdown("### 💬 Ask Questions")
    gr.ChatInterface(fn=_chat_bridge, title="Q&A Interface", type="messages")

if __name__ == "__main__":
    # pip install gradio networkx python-docx PyPDF2 python-dotenv community matplotlib pandas
    print("Starting Gradio app...")
    print(f"Demo object: {demo}")
    demo.launch(show_api=False, share=True)
